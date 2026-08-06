"""
Extract plain text from .docx files (paragraphs and table cells).

Used by WordValidatorService before required-string checks.
"""

import logging
import re
from pathlib import Path

logger = logging.getLogger(__name__)


class WordReaderService:
    """Loads documents with python-docx and builds a single searchable text blob."""

    def extract_text(self, file_path: Path) -> tuple[str, int]:
        """
        Read body text and count paragraphs.

        Returns:
            (combined_text, paragraph_count)

        Raises:
            FileNotFoundError: path is not a file.
            RuntimeError: python-docx is not installed.
        """
        if not file_path.is_file():
            raise FileNotFoundError(f"Word document not found: {file_path}")

        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError(
                "python-docx is not installed. Add python-docx to requirements.txt."
            ) from exc

        document = Document(str(file_path))
        parts: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        parts.append(text)

        body = "\n".join(parts)
        logger.info(
            "Read Word document %s (%s paragraphs, %s characters)",
            file_path.name,
            len(document.paragraphs),
            len(body),
        )
        return body, len(document.paragraphs)

    @staticmethod
    def normalize(text: str) -> str:
        """Lowercase and collapse whitespace for substring matching."""
        return re.sub(r"\s+", " ", text or "").strip().lower()
