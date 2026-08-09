#!/usr/bin/env python3
"""
Build Windows-Server-Deployment-Guide.docx from the markdown source.

Usage (from repo root, using backend venv):

  cd backend
  .venv\\Scripts\\activate          # Windows
  # source .venv/bin/activate         # macOS/Linux
  pip install python-docx matplotlib
  set MPLCONFIGDIR=..\\docs\\.matplotlib   # Windows CMD if matplotlib cache errors
  python ..\\docs\\build_deployment_guide_docx.py

Output: docs/Windows-Server-Deployment-Guide.docx
"""

from __future__ import annotations

import re
import tempfile
from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch

DOCS_DIR = Path(__file__).resolve().parent
MD_PATH = DOCS_DIR / "Windows-Server-Deployment-Guide.md"
OUT_PATH = DOCS_DIR / "Windows-Server-Deployment-Guide.docx"

def _render_architecture_diagram_png(path: Path) -> None:
    """Draw deployment architecture (matplotlib, no external Mermaid CLI)."""
    fig, ax = plt.subplots(figsize=(12, 5))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 5)
    ax.axis("off")

    def box(x, y, w, h, text, fc="#E8F4FC", ec="#2E75B6"):
        patch = FancyBboxPatch(
            (x, y),
            w,
            h,
            boxstyle="round,pad=0.02,rounding_size=0.08",
            linewidth=1.2,
            edgecolor=ec,
            facecolor=fc,
        )
        ax.add_patch(patch)
        ax.text(
            x + w / 2,
            y + h / 2,
            text,
            ha="center",
            va="center",
            fontsize=8,
            wrap=True,
        )

    def cylinder(x, y, w, h, text):
        patch = FancyBboxPatch(
            (x, y),
            w,
            h,
            boxstyle="round,pad=0.02,rounding_size=0.05",
            linewidth=1.2,
            edgecolor="#548235",
            facecolor="#E2EFDA",
        )
        ax.add_patch(patch)
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=7)

    def arrow(x1, y1, x2, y2):
        ax.add_patch(
            FancyArrowPatch(
                (x1, y1),
                (x2, y2),
                arrowstyle="-|>",
                mutation_scale=12,
                linewidth=1,
                color="#404040",
            )
        )

    # Row 1: user -> IIS -> split
    box(0.2, 2.0, 1.4, 0.9, "User\nBrowser")
    box(2.0, 2.0, 1.6, 0.9, "React UI\n(frontend/dist)", fc="#FFF2CC", ec="#BF9000")
    box(4.0, 2.0, 1.8, 0.9, "IIS\nHTTPS + Rewrite", fc="#FCE4D6", ec="#C65911")
    box(6.2, 2.8, 2.0, 0.85, "Static files\n(SPA)", fc="#FFF2CC", ec="#BF9000")
    box(6.2, 1.2, 2.0, 0.85, "FastAPI\nuvicorn :8000", fc="#DDEBF7", ec="#2E75B6")

    arrow(1.6, 2.45, 2.0, 2.45)
    arrow(3.6, 2.45, 4.0, 2.45)
    arrow(5.8, 2.55, 6.2, 2.9)
    arrow(5.8, 2.35, 6.2, 1.65)

    # API services
    box(8.6, 0.3, 1.5, 0.7, "Excel\nprocessor")
    box(8.6, 1.2, 1.5, 0.7, "PDF\nvalidator")
    box(8.6, 2.1, 1.5, 0.7, "Word\nvalidator")
    box(8.6, 3.0, 1.5, 0.7, "Upload /\nPath")

    arrow(8.2, 1.6, 8.6, 0.65)
    arrow(8.2, 1.6, 8.6, 1.55)
    arrow(8.2, 1.6, 8.6, 2.45)
    arrow(8.2, 1.6, 8.6, 3.35)

    cylinder(10.4, 0.25, 1.3, 0.75, "storage/\nuploads")
    cylinder(10.4, 1.15, 1.3, 0.75, "storage/\nprocessed")
    box(10.4, 2.2, 1.3, 0.7, "PASS/\nFAIL", fc="#E2EFDA", ec="#548235")
    box(10.4, 3.05, 1.3, 0.7, "Optional\nin-place", fc="#F2F2F2", ec="#7F7F7F")

    arrow(10.1, 0.65, 10.4, 0.65)
    arrow(10.1, 1.55, 10.4, 1.55)
    arrow(10.1, 2.45, 10.4, 2.55)
    arrow(10.1, 3.35, 10.4, 3.4)

    ax.set_title(
        "Production architecture — Windows Server + IIS + domain (HTTPS)",
        fontsize=11,
        fontweight="bold",
        pad=12,
    )
    fig.tight_layout()
    fig.savefig(path, dpi=160, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def _add_formatted_paragraph(doc: Document, line: str, style: str | None = None):
    """Paragraph with **bold** segments."""
    p = doc.add_paragraph(style=style)
    parts = re.split(r"(\*\*[^*]+\*\*)", line)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)
    return p


def _parse_table_rows(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        if not line.strip().startswith("|"):
            break
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if all(re.match(r"^[-:\s]+$", c) for c in cells):
            continue
        rows.append(cells)
    return rows


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(cols):
            text = row[ci] if ci < len(row) else ""
            text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
            text = text.replace("`", "")
            cell = table.rows[ri].cells[ci]
            cell.text = text
            if ri == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def build_docx() -> None:
    if not MD_PATH.is_file():
        raise FileNotFoundError(MD_PATH)

    md_text = MD_PATH.read_text(encoding="utf-8")
    lines = md_text.splitlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Windows Server Deployment Guide", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub = doc.add_paragraph("Document Processing POC")
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub.runs[0].italic = True

    doc.add_heading("Architecture diagram", level=1)
    doc.add_paragraph(
        "Browser traffic hits IIS on your domain (HTTPS). IIS serves the React build "
        "and reverse-proxies API routes to FastAPI on localhost port 8000."
    )
    with tempfile.TemporaryDirectory() as tmp:
        png = Path(tmp) / "architecture.png"
        _render_architecture_diagram_png(png)
        doc.add_picture(str(png), width=Inches(6.5))

    doc.add_paragraph()
    cap = doc.add_paragraph("Figure 1 — Deployment architecture (single-domain layout)")
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(10)
    cap.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    i = 0
    in_code = False
    code_lang = ""
    code_lines: list[str] = []
    table_buffer: list[str] = []

    def flush_table() -> None:
        nonlocal table_buffer
        if table_buffer:
            _add_table(doc, _parse_table_rows(table_buffer))
            table_buffer = []
            doc.add_paragraph()

    while i < len(lines):
        line = lines[i]

        if line.strip() == "---":
            flush_table()
            i += 1
            continue

        if line.startswith("```"):
            flush_table()
            if not in_code:
                in_code = True
                code_lang = line.strip("`").strip() or "text"
                code_lines = []
            else:
                in_code = False
                p = doc.add_paragraph()
                p.style = "No Spacing"
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.25)
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if line.strip().startswith("|"):
            table_buffer.append(line)
            i += 1
            continue
        flush_table()

        if line.startswith("# ") and not line.startswith("## "):
            i += 1
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue

        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("- [ ]"):
            doc.add_paragraph(stripped[5:].strip(), style="List Bullet")
            i += 1
            continue

        if stripped.startswith("- "):
            _add_formatted_paragraph(doc, stripped[2:], style="List Bullet")
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            _add_formatted_paragraph(doc, text, style="List Number")
            i += 1
            continue

        if stripped.startswith("*") and stripped.endswith("*"):
            p = doc.add_paragraph(stripped.strip("*"))
            p.runs[0].italic = True
            i += 1
            continue

        _add_formatted_paragraph(doc, stripped)
        i += 1

    flush_table()

    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    build_docx()
